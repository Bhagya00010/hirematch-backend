from __future__ import annotations

import hashlib
import logging
import math
import re
from datetime import datetime
from pathlib import Path
from typing import Any
from uuid import UUID, uuid4

from fastapi import UploadFile
from sqlalchemy import select, text
from sqlalchemy.orm import Session, joinedload

from app.core.config import settings
from app.core.llm import get_embeddings, get_llm
from app.models.job import Job, JobEmbedding
from app.models.resume import (
    Candidate,
    CandidateEmbedding,
    MatchResult,
    ResumeFile,
    ResumeProcessingStatus,
    ResumeValidationStatus,
)
from app.services.workflow import parse_json_response

logger = logging.getLogger(__name__)

ALLOWED_RESUME_EXTENSIONS = {".pdf", ".docx", ".txt"}

SKILL_SYNONYMS: dict[str, str] = {
    "aws ec2": "aws", "aws s3": "aws", "aws lambda": "aws", "aws rds": "aws",
    "aws sqs": "aws", "aws sns": "aws", "aws dynamodb": "aws",
    "amazon web services": "aws", "amazon ec2": "aws", "amazon s3": "aws",
    "apache kafka": "kafka", "kafka streams": "kafka",
    "ci/cd": "ci/cd pipelines", "continuous integration": "ci/cd pipelines",
    "continuous deployment": "ci/cd pipelines", "cicd": "ci/cd pipelines",
    "gitlab ci": "ci/cd pipelines", "github actions": "ci/cd pipelines",
    "jenkins": "ci/cd pipelines",
    "rag": "retrieval augmented generation",
    "retrieval-augmented generation": "retrieval augmented generation",
    "vector db": "vector databases", "vector database": "vector databases",
    "llm orchestration": "llm orchestration frameworks",
    "microservices": "distributed systems", "microservice": "distributed systems",
    "microservices architecture": "distributed systems",
    "full stack": "full stack architecture", "fullstack": "full stack architecture",
    "full-stack": "full stack architecture",
    "reactjs": "react", "react.js": "react",
    "angularjs": "angular", "angular.js": "angular",
    "vuejs": "vue", "vue.js": "vue",
    "nodejs": "node.js",
    "machine learning": "ml", "artificial intelligence": "ai",
    "deep learning": "dl", "neural networks": "deep learning",
    "docker container": "docker", "containerization": "docker",
    "k8s": "kubernetes", "k8": "kubernetes",
    "kubernetes orchestration": "kubernetes",
    "rest api": "rest", "restful": "rest",
    "gql": "graphql",
    "nosql": "mongodb", "relational database": "sql", "rdbms": "sql",
    "postgres": "postgresql", "mongo": "mongodb",
    "redis cache": "redis", "elastic search": "elasticsearch",
    "springboot": "spring boot", "spring-boot": "spring boot",
    "spark": "apache spark",
    "model deployment": "mlops", "model ops": "mlops",
    "genai": "generative ai", "agentic ai": "agentic ai solutions",
    "data structures & algorithms": "data structures",
    "dsa": "data structures",
    "oop": "oop concepts",
    "object oriented programming": "oop concepts",
    "enterprise architecture": "enterprise application architecture",
    "batch processing": "data processing",
    "real-time processing": "data processing",
    "typescript": "typescript", "javascript": "javascript",
    # FIX P1: architecture family synonyms
    "solution architecture": "enterprise application architecture",
    "technical architecture": "enterprise application architecture",
    "software architecture": "enterprise application architecture",
    "system architecture": "enterprise application architecture",
    "microservice architecture": "distributed systems",
    "serverless architecture": "enterprise application architecture",
    "monorepo": "enterprise application architecture",
    "event driven architecture": "distributed systems",
    "event-driven architecture": "distributed systems",
    "eda": "distributed systems",
}

SKILL_ONTOLOGY: dict[str, list[str]] = {
    "java": [
        "oop concepts", "oop", "object oriented programming",
        "multithreading", "collections", "data structures", "jvm",
        "algorithms", "design patterns", "solid principles",
    ],
    "kotlin": [
        "oop concepts", "jvm", "android",
        "data structures", "algorithms", "design patterns",
    ],
    "scala": [
        "oop concepts", "jvm", "functional programming",
        "data structures", "algorithms",
    ],
    "c#": [
        "oop concepts", "object oriented programming", "data structures",
        "algorithms", "design patterns", ".net",
    ],
    "python": [
        "scripting", "data structures", "oop concepts", "algorithms",
    ],
    "go": [
        "data structures", "algorithms", "concurrency", "microservices",
    ],
    "c++": [
        "oop concepts", "data structures", "algorithms",
        "systems programming", "memory management",
    ],
    "kotlin": ["oop concepts", "jvm", "android", "data structures", "algorithms"],
    "apache spark": ["distributed systems", "data processing", "big data",
                     "batch processing", "real-time processing"],
    "kafka": ["distributed systems", "event driven", "data processing",
              "messaging", "pub sub"],
    "flink": ["distributed systems", "real-time processing", "data processing"],
    "hadoop": ["distributed systems", "data processing", "big data", "hdfs"],
    "spring boot": [
        "microservices", "rest", "distributed systems",
        "dependency injection", "spring framework",
        "enterprise application architecture",  # NEW
    ],
    "spring framework": ["dependency injection", "oop concepts"],
    "aws": ["cloud computing", "cloud infrastructure", "s3", "ec2"],
    "azure": ["cloud computing", "cloud infrastructure"],
    "gcp": ["cloud computing", "cloud infrastructure"],
    "kubernetes": ["container orchestration", "docker", "devops",
                   "distributed systems"],
    "docker": ["containerization", "devops"],
    "pytorch": ["deep learning", "machine learning", "neural networks", "python"],
    "tensorflow": ["deep learning", "machine learning", "neural networks", "python"],
    "scikit-learn": ["machine learning", "python", "data science"],
    "fastapi": ["rest", "python", "microservices"],
    "django": ["rest", "python", "web development", "oop concepts"],
    "react": ["javascript", "frontend", "web development", "ui development"],
    "angular": ["javascript", "typescript", "frontend", "web development"],
    "vue": ["javascript", "frontend", "web development"],
    "postgresql": ["sql", "rdbms", "relational database"],
    "mysql": ["sql", "rdbms", "relational database"],
    "mongodb": ["nosql", "database"],
    "elasticsearch": ["nosql", "search", "data indexing"],
    "redis": ["caching", "in-memory database", "nosql"],
    "distributed systems": ["system design", "scalability", "high availability"],
    "microservices": [
        "distributed systems", "system design", "rest",
        "enterprise application architecture",   # FIX P1
    ],
    "serverless": [
        "cloud computing", "enterprise application architecture",
        "distributed systems",
    ],
    "system design": ["scalability", "high availability", "architecture"],
    "mlops": ["model deployment", "ci/cd pipelines", "devops", "machine learning"],
    "ci/cd pipelines": ["devops", "automation"],
    "dbt": ["data engineering", "sql", "data transformation"],
    "airflow": ["data engineering", "workflow orchestration", "python"],
    "snowflake": ["data warehouse", "sql", "cloud computing"],
    "enterprise application architecture": [
        "system design", "distributed systems", "scalability",
        "high availability", "microservices",
    ],
}

_ONTOLOGY_REVERSE: dict[str, list[str]] = {}
for _parent, _children in SKILL_ONTOLOGY.items():
    for _child in _children:
        _ONTOLOGY_REVERSE.setdefault(_child.lower(), []).append(_parent.lower())

ARCHITECTURE_FAMILY_SKILLS: set[str] = {
    "microservices", "distributed systems", "serverless", "serverless architecture",
    "microservice architecture", "event driven architecture", "event-driven architecture",
    "spring boot", "spring cloud", "monorepo", "system design",
    "high level design", "hld", "solution architecture", "technical architecture",
    "software architecture", "api gateway", "service mesh",
    "domain driven design", "ddd", "hexagonal architecture",
    "clean architecture", "cqrs", "event sourcing",
}


_ARCH_FAMILY_CONFIDENCE = 0.72   # above soft-match floor (0.60), below mandatory gate (0.75)


SKILL_CLUSTERS: dict[str, set[str]] = {
    "data_structures_algorithms": {
        "data structures", "algorithms", "dsa",
        "data structures & algorithms", "data structures and algorithms",
        "algo", "algorithmic thinking",
    },
    "oop": {
        "oop", "oop concepts", "object oriented programming",
        "object-oriented programming", "oops", "oops concepts",
    },
    "system_design": {
        "system design", "systems design", "high level design",
        "low level design", "hld", "lld",
    },
    "enterprise_architecture": {
        "enterprise architecture", "enterprise application architecture",
        "solution architecture", "technical architecture",
        "software architecture", "system architecture",
    },
    "distributed_systems": {
        "distributed systems", "distributed computing",
        "large scale systems", "scalability",
    },
    "ci_cd": {
        "ci/cd", "ci/cd pipelines", "continuous integration",
        "continuous deployment", "continuous delivery",
        "github actions", "gitlab ci", "jenkins",
    },
    "cloud_aws": {
        "aws", "amazon web services", "aws ec2", "aws s3",
        "aws lambda", "aws rds", "aws sqs",
    },
}

_CLUSTER_REVERSE: dict[str, str] = {}
for _cluster, _aliases in SKILL_CLUSTERS.items():
    for _alias in _aliases:
        _CLUSTER_REVERSE[_alias.lower()] = _cluster


def get_cluster_key(skill: str) -> str:
    """Return cluster key if skill belongs to a cluster, else the skill itself."""
    return _CLUSTER_REVERSE.get(skill.lower().strip(), skill.lower().strip())


def deduplicate_skills_by_cluster(skills: list[str]) -> list[str]:
    """
    Remove duplicate skills that belong to the same cluster.
    Returns one representative per cluster, keeps all non-clustered skills.
    """
    seen_clusters: set[str] = set()
    out: list[str] = []
    for s in skills:
        ck = get_cluster_key(s)
        if ck in seen_clusters:
            continue
        seen_clusters.add(ck)
        out.append(s)
    return out

# Stop-words

STOP_WORDS: set[str] = {
    "role", "developer", "general", "stack", "certified", "certification",
    "responsibility", "candidate", "job", "position", "work", "experience",
    "years", "year", "team", "company", "business", "organization",
    "looking", "seeking", "hiring", "recruiting", "apply", "application",
    "full", "web", "end", "end-to", "end-to-end",
    "untitled", "development", "preferred", "nice to have", "good to have",
    "bonus", "plus",
}

RESUME_KEYWORDS: set[str] = {
    "resume", "curriculum vitae", "experience", "education", "skills",
    "key skills", "professional summary", "professional experience",
    "work experience", "projects", "technologies", "employment", "certifications",
}

FINANCIAL_DOC_PATTERNS = [
    r"total amount due", r"bill to\s*:", r"invoice\s*#\s*\d+",
    r"invoice\s+number\s*:", r"purchase order\s*#", r"tax invoice",
    r"payment due\s*(date|by)", r"subtotal\s*[:\$]", r"amount payable",
]
RESUME_PROJECT_CONTEXT = [
    r"(built|developed|created|designed|managed|automated|integrated|processed)\b.{0,60}\binvoice",
    r"\binvoice.{0,60}(system|module|portal|platform|application|software|tool|feature|api|service)",
    r"(project|experience|work).{0,200}\binvoice",
]



# Section 1  –  Skill normalisation

def apply_synonym_normalization(skill: str) -> str:
    if not skill:
        return skill
    key = skill.lower().strip()
    if key in SKILL_SYNONYMS:
        return SKILL_SYNONYMS[key]
    for k, v in SKILL_SYNONYMS.items():
        if v.lower() == key:
            return v
    return skill


def normalize_keywords(values: list[str | None]) -> set[str]:
    out: set[str] = set()
    for v in values:
        if not v:
            continue
        mapped = apply_synonym_normalization(str(v)).lower().strip()
        if mapped and mapped not in STOP_WORDS:
            out.add(mapped)
    return out


normalize_keywords_with_synonyms = normalize_keywords


_must_have_weight_cache: dict[str, dict[str, str]] = {}

_TIER_WEIGHT = {"mandatory": 1.0, "preferred": 0.6, "bonus": 0.3}
_TIER_COMPAT = {
    "critical": "mandatory", "important": "preferred", "optional": "bonus",
    "mandatory": "mandatory", "preferred": "preferred", "bonus": "bonus",
}


def get_weighted_must_have_skills(job: Job) -> dict[str, str]:
    """
    FIX P7 (v5): Rule-based mandatory tier assignment.

    MANDATORY (rule-based, no LLM, cannot be downgraded):
      - job.required_skills      → recruiter explicitly typed these in.
      - job.ai_required_skills   → AI parser extracted these from the JD's
                                    "Required:" / "Must Have:" section. These
                                    represent skills the JD ITSELF marks as
                                    required, so they are treated with the
                                    same authority as recruiter_required.

    PREFERRED / BONUS (LLM-classified, secondary pools only):
      - job.ai_must_have_keywords
      - job.ai_nice_to_have_keywords
      - job.ai_tools
      - job.ai_technologies

    Anything already counted as mandatory is removed from the secondary pool
    before LLM classification (no double counting).

    Returns  { skill_name_lower: "mandatory" | "preferred" | "bonus" }
    """
    cache_key = str(job.job_id)
    if cache_key in _must_have_weight_cache:
        return _must_have_weight_cache[cache_key]

    recruiter_required: set[str] = set(
        s.lower().strip() for s in (job.required_skills or []) if s and s.strip()
    )
    
    ai_required: set[str] = set(
        s.lower().strip() for s in (job.ai_required_skills or []) if s and s.strip()
    )

    mandatory_raw = recruiter_required | ai_required

    mandatory_deduped = deduplicate_skills_by_cluster(sorted(mandatory_raw))
    result: dict[str, str] = {s: "mandatory" for s in mandatory_deduped}

    secondary_skills: list[str] = list(set(
        (job.ai_must_have_keywords or [])
        + (job.ai_nice_to_have_keywords or [])
        + (job.ai_tools or [])
        + (job.ai_technologies or [])
    ))
    
    secondary_skills = [
        s for s in secondary_skills
        if s.lower().strip() not in mandatory_raw
        and get_cluster_key(s) not in {get_cluster_key(m) for m in mandatory_deduped}
    ]

    secondary_deduped = deduplicate_skills_by_cluster(secondary_skills)

    print(f"\n[MUST_HAVE_RULES] Job='{job.job_title}'")
    print(f"[MUST_HAVE_RULES] recruiter_required ({len(recruiter_required)}): {sorted(recruiter_required)}")
    print(f"[MUST_HAVE_RULES] ai_required_skills ({len(ai_required)}): {sorted(ai_required)}")
    print(f"[MUST_HAVE_RULES] => MANDATORY total ({len(mandatory_deduped)}): {mandatory_deduped}")
    print(f"[MUST_HAVE_RULES] Secondary pool for LLM classification ({len(secondary_deduped)}): {secondary_deduped}")

    if secondary_deduped:
        prompt = f"""You are a senior technical recruiter. Classify each skill into:

- preferred : strongly desired but NOT a dealbreaker (recruiter did not mark it as "Required")
- bonus     : nice-to-have / peripheral

IMPORTANT: Do NOT classify anything as "mandatory". That tier is reserved for skills the JD
explicitly lists under "Required". Your options are ONLY "preferred" or "bonus".

Job title       : {job.job_title}
Job description : {(job.job_description or '')[:1500]}

Skills to classify:
{chr(10).join(f'- {s}' for s in secondary_deduped)}

Return ONLY valid JSON. No markdown, no extra keys.
Format: {{"<skill_name_lower>": "<preferred|bonus>", ...}}
"""
        try:
            llm = get_llm(temperature=0.0)
            response = llm.invoke(prompt)
            raw = response.content if hasattr(response, "content") else str(response)
            print(f"[MUST_HAVE_LLM] Raw: {raw[:300]}")
            parsed: dict = parse_json_response(raw)
            for k, v in parsed.items():
                sk = k.lower().strip()
                if sk in result:  # already mandatory via rule
                    continue
                raw_tier = str(v).lower()
                if raw_tier == "mandatory":
                    raw_tier = "preferred"
                    print(f"[MUST_HAVE_LLM] CAP: '{sk}' downgraded mandatory→preferred (secondary pool)")
                normalized_tier = _TIER_COMPAT.get(raw_tier, "preferred")
                result[sk] = normalized_tier
            # Default unclassified secondary skills to preferred
            for s in secondary_deduped:
                sk = s.lower().strip()
                if sk not in result:
                    result[sk] = "preferred"
        except Exception as exc:
            print(f"[MUST_HAVE_LLM] ERROR: {exc} – defaulting secondary skills to 'preferred'")
            logger.warning("Skill tier classification failed: %s", exc)
            for s in secondary_deduped:
                sk = s.lower().strip()
                if sk not in result:
                    result[sk] = "preferred"

    _must_have_weight_cache[cache_key] = result
    print(f"[MUST_HAVE_RULES] Final ({len(result)}): {result}")
    logger.info("Skill tier classification for job %s: %s", job.job_id, result)
    return result


# Section 3  –  Confidence-gated skill matching


_MATCH_CONFIDENCE: dict[str, float] = {
    "exact":         1.00,
    "boundary":      0.95,
    "fuzzy":         0.85,
    "ontology":      0.78,
    "arch_family":   0.72,   
    "semantic":      0.70,   
}
_MIN_CONFIDENCE = 0.75   


def _direct_match(target_lower: str, cs_lower: str) -> bool:
    if target_lower == cs_lower:
        return True
    if re.search(r"\b" + re.escape(target_lower) + r"\b", cs_lower):
        return True
    if re.search(r"\b" + re.escape(cs_lower) + r"\b", target_lower):
        return True
    return False


def _is_architecture_target(target_lower: str) -> bool:
    """Return True if target is an enterprise/solution architecture skill."""
    arch_targets = {
        "enterprise application architecture", "enterprise architecture",
        "solution architecture", "technical architecture",
        "software architecture", "system architecture",
    }
    return target_lower in arch_targets


def _candidate_has_architecture_family(candidate_skills_lower: list[str]) -> bool:
    """
    FIX P1: Return True if candidate has 2+ architecture-family skills,
    indicating real architecture experience even without the exact label.
    """
    count = sum(
        1 for cs in candidate_skills_lower
        if cs in ARCHITECTURE_FAMILY_SKILLS
        or any(_direct_match(af, cs) for af in ARCHITECTURE_FAMILY_SKILLS)
    )
    return count >= 2


def _skill_match_with_confidence(
    target: str,
    candidate_skills: list[str],
) -> tuple[bool, float, str]:
    """
    Return (matched, confidence, match_type).

    FIX P1: Architecture family match at 0.72 confidence (partial credit).
    FIX P2: Ontology now includes Java→DSA, Java→OOP, etc.

    Match strategies in priority order:
    1. Exact / word-boundary       → confidence 0.95–1.0
    2. Fuzzy token_set_ratio ≥ 85  → confidence 0.85
    3. Ontology expansion          → confidence 0.78
    4. Architecture family         → confidence 0.72 (FIX P1 – partial credit only)
    """
    target_lower = target.lower().strip()
    candidate_lowers = [cs.lower().strip() for cs in candidate_skills]

    for cs_lower in candidate_lowers:
        if target_lower == cs_lower:
            return True, _MATCH_CONFIDENCE["exact"], "exact"
        if re.search(r"\b" + re.escape(target_lower) + r"\b", cs_lower):
            return True, _MATCH_CONFIDENCE["boundary"], "boundary"
        if re.search(r"\b" + re.escape(cs_lower) + r"\b", target_lower):
            return True, _MATCH_CONFIDENCE["boundary"], "boundary"

    try:
        from thefuzz import fuzz
        for cs_lower in candidate_lowers:
            ratio = fuzz.token_set_ratio(target_lower, cs_lower)
            if ratio >= 85:
                confidence = _MATCH_CONFIDENCE["fuzzy"] * (ratio / 100)
                if confidence >= _MIN_CONFIDENCE:
                    return True, confidence, "fuzzy"
    except ImportError:
        for cs_lower in candidate_lowers:
            ratio = _simple_token_ratio(target_lower, cs_lower)
            if ratio >= 85:
                confidence = _MATCH_CONFIDENCE["fuzzy"] * (ratio / 100)
                if confidence >= _MIN_CONFIDENCE:
                    return True, confidence, "fuzzy"

    ontology_conf = _MATCH_CONFIDENCE["ontology"]
    if ontology_conf >= _MIN_CONFIDENCE:
        parent_skills = _ONTOLOGY_REVERSE.get(target_lower, [])
        for parent in parent_skills:
            for cs_lower in candidate_lowers:
                if _direct_match(parent, cs_lower):
                    return True, ontology_conf, "ontology"
        for cs_lower in candidate_lowers:
            implied = SKILL_ONTOLOGY.get(cs_lower, [])
            if target_lower in [i.lower() for i in implied]:
                return True, ontology_conf, "ontology"

    if _is_architecture_target(target_lower):
        if _candidate_has_architecture_family(candidate_lowers):
            arch_conf = _MATCH_CONFIDENCE["arch_family"]
            print(f"  [ARCH_FAMILY] '{target_lower}' matched via architecture-family cluster (conf={arch_conf})")
            return True, arch_conf, "arch_family"

    return False, 0.0, "none"


def _skill_matched(target: str, candidate_skills: list[str]) -> bool:
    """Backward-compatible wrapper – returns bool only (uses >= _MIN_CONFIDENCE)."""
    matched, confidence, _ = _skill_match_with_confidence(target, candidate_skills)
    return matched and confidence >= _MIN_CONFIDENCE


# Section 4  –  Must-have scoring with cluster dedup + confidence gate

def recency_weight(last_used_year: int | None) -> float:
    if last_used_year is None:
        return 0.8
    gap = datetime.now().year - int(last_used_year)
    if gap <= 1:  return 1.0
    if gap <= 2:  return 0.92
    if gap <= 3:  return 0.85
    if gap <= 5:  return 0.65
    if gap <= 8:  return 0.45
    return 0.3


def build_skill_proficiency_map(parsed: dict) -> dict[str, dict]:
    proficiency_map: dict[str, dict] = {}
    for entry in (parsed.get("skill_proficiency") or []):
        if not isinstance(entry, dict):
            continue
        skill = str(entry.get("skill", "")).lower().strip()
        if skill:
            proficiency_map[skill] = {
                "years": entry.get("years"),
                "last_used_year": entry.get("last_used_year"),
                "confidence": entry.get("confidence", 0.8),
            }
    return proficiency_map


def score_must_have_skills(
    job: Job,
    candidate: Candidate,
    weighted_skills: dict[str, str],
    skill_proficiency_map: dict[str, dict] | None = None,
) -> tuple[float, list[str]]:
    """
    FIX P1 + P2 + P3: Weighted must-have score with:
    - Confidence gating: mandatory gate requires >= 0.75
    - Architecture family: contributes at 0.72 confidence (partial score)
    - Cluster deduplication: DSA/data structures/algorithms → counted once
    - Recency weighting

    Returns (score_0_to_100, missing_mandatory_skills_list)
    """
    if not weighted_skills:
        return 0.0, []

    skill_proficiency_map = skill_proficiency_map or {}

    deduped_skills: dict[str, str] = {}
    seen_clusters: set[str] = set()
    for skill, tier in weighted_skills.items():
        ck = get_cluster_key(skill)
        if ck in seen_clusters:
            print(f"  [MUST_HAVE_SCORE] CLUSTER DEDUP: skipping '{skill}' (cluster '{ck}' already counted)")
            continue
        seen_clusters.add(ck)
        deduped_skills[skill] = tier

    candidate_skills: list[str] = (candidate.skills or []) + (candidate.tech_stack or [])
    print(f"  [MUST_HAVE_SCORE] Candidate='{getattr(candidate,'full_name','?')}' | candidate_skills={candidate_skills[:10]}")
    print(f"  [MUST_HAVE_SCORE] Skills after cluster dedup ({len(deduped_skills)}): {dict(list(deduped_skills.items())[:10])}")

    total_possible = sum(_TIER_WEIGHT.get(v, 0.6) for v in deduped_skills.values())
    if total_possible == 0:
        return 0.0, []

    earned = 0.0
    missing_mandatory: list[str] = []

    for skill, tier in deduped_skills.items():
        tier_w = _TIER_WEIGHT.get(tier, 0.6)
        matched_flag, confidence, match_type = _skill_match_with_confidence(skill, candidate_skills)

        if matched_flag:
            prof = skill_proficiency_map.get(skill.lower().strip(), {})
            rec_w = recency_weight(prof.get("last_used_year"))

            # FIX P1: arch_family match contributes proportionally to confidence
            # 0.72 confidence → 72% of full tier weight (partial credit)
            conf_multiplier = confidence if confidence < _MIN_CONFIDENCE else 1.0
            contribution = tier_w * rec_w * conf_multiplier
            earned += contribution
            print(
                f"  [MUST_HAVE_SCORE]   ✓ '{skill}' ({tier}, conf={confidence:.2f}, "
                f"type={match_type}, recency={rec_w:.2f}, conf_mult={conf_multiplier:.2f}) "
                f"earned+={contribution:.2f}"
            )
            # For mandatory gate: arch_family (0.72) does NOT count as matched
            # (gate uses _skill_matched which requires >= _MIN_CONFIDENCE = 0.75)
            if tier == "mandatory" and confidence < _MIN_CONFIDENCE:
                missing_mandatory.append(f"{skill} (partial:{match_type})")
        else:
            print(f"  [MUST_HAVE_SCORE]   ✗ '{skill}' ({tier}) NOT matched (conf=0)")
            if tier == "mandatory":
                missing_mandatory.append(skill)

    score = round((earned / total_possible) * 100, 2)
    print(f"  [MUST_HAVE_SCORE] earned={earned:.2f}/possible={total_possible:.2f} → score={score}  missing_mandatory={missing_mandatory}")
    return score, missing_mandatory


# Section 5  –  Mandatory Skill Gate

def apply_mandatory_gate(
    candidate: Candidate,
    weighted_skills: dict[str, str],
    idx: int,
) -> tuple[bool, float]:
    """
    Gate fires only when MANDATORY tier skills are < 40% matched.
    Uses _skill_matched (requires confidence >= 0.75).

    FIX P1: arch_family confidence (0.72) does NOT satisfy mandatory gate.
    FIX P7: mandatory skills are now rule-based (recruiter required_skills ∪
            ai_required_skills) — see get_weighted_must_have_skills().
    Returns (gate_failed: bool, critical_ratio: float).
    """
    mandatory_skills = [s for s, t in weighted_skills.items() if t == "mandatory"]

    if not mandatory_skills:
        print(f"[GATE #{idx}] No mandatory skills defined – gate skipped")
        return False, 1.0

    candidate_skills_flat = (candidate.skills or []) + (candidate.tech_stack or [])
    matched_count = sum(
        1 for s in mandatory_skills
        if _skill_matched(s, candidate_skills_flat)
    )
    ratio = matched_count / len(mandatory_skills)
    print(f"[GATE #{idx}] mandatory={len(mandatory_skills)} matched={matched_count} ratio={ratio:.2f}")

    if ratio < 0.40:
        print(f"[GATE #{idx}] ⚠ MANDATORY GATE FAILED – will apply ×0.20 penalty")
        return True, ratio

    return False, ratio


# Section 6  –  FIX P6: Project Impact Scoring

_project_score_cache: dict[str, float] = {}


def score_project_impact(job: Job, candidate: Candidate) -> float:
    """
    FIX P6: Score candidate's projects on 4 dimensions:
    - Complexity   : technical difficulty, architecture depth
    - Scale        : users served, data volume, throughput
    - Ownership    : individual contribution vs team size
    - Business Impact: revenue, cost, user growth, strategic value

    Returns 0-100 score. Cached per candidate_id.
    """
    cache_key = str(candidate.id)
    if cache_key in _project_score_cache:
        return _project_score_cache[cache_key]

    raw_text = (candidate.raw_text or "")[:4000]
    if not raw_text:
        _project_score_cache[cache_key] = 40.0  # neutral default
        return 40.0

    prompt = f"""You are a senior engineering manager evaluating a candidate's project impact.

Score this candidate's projects/experience across 4 dimensions (0-100 each):

1. complexity      : How technically complex? (simple CRUD=20, microservices=60, distributed at scale=90+)
2. scale           : How large? (personal project=20, 1K users=40, 100K users=70, 10M+ users=90+)
3. ownership       : Individual ownership? (small part of large team=20, lead contributor=70, sole owner=90+)
4. business_impact : Real business outcome? (internal tool=20, improved metrics=60, revenue/growth impact=90+)

Job being applied to: {job.job_title}
Job description (brief): {(job.job_description or '')[:500]}

Candidate's experience/projects:
{raw_text}

Return ONLY valid JSON:
{{
  "complexity": <0-100>,
  "scale": <0-100>,
  "ownership": <0-100>,
  "business_impact": <0-100>,
  "top_project_summary": "<one sentence describing their strongest project>"
}}
"""
    print(f"  [PROJECT_SCORE] Scoring projects for: {candidate.full_name or 'Unknown'}")
    try:
        llm = get_llm(temperature=0.0)
        response = llm.invoke(prompt)
        raw = response.content if hasattr(response, "content") else str(response)
        parsed = parse_json_response(raw)

        complexity     = max(0.0, min(100.0, float(parsed.get("complexity", 40))))
        scale          = max(0.0, min(100.0, float(parsed.get("scale", 40))))
        ownership      = max(0.0, min(100.0, float(parsed.get("ownership", 40))))
        business_impact = max(0.0, min(100.0, float(parsed.get("business_impact", 40))))

        # Weighted blend: complexity matters most for technical roles
        project_score = round(
            complexity     * 0.30
            + scale        * 0.25
            + ownership    * 0.25
            + business_impact * 0.20,
            2,
        )
        print(
            f"  [PROJECT_SCORE] complexity={complexity:.0f} scale={scale:.0f} "
            f"ownership={ownership:.0f} impact={business_impact:.0f} → {project_score:.1f}"
        )
        print(f"  [PROJECT_SCORE] Top project: {parsed.get('top_project_summary', 'N/A')}")
        logger.info(
            "Project score %s: %.1f (cplx=%.0f scale=%.0f own=%.0f impact=%.0f)",
            candidate.full_name, project_score, complexity, scale, ownership, business_impact,
        )
        _project_score_cache[cache_key] = project_score
        return project_score
    except Exception as exc:
        print(f"  [PROJECT_SCORE] ERROR: {exc} – using default 40.0")
        logger.warning("Project score failed for %s: %s", candidate.full_name, exc)
        _project_score_cache[cache_key] = 40.0
        return 40.0


# Section 7  –  LangChain EnsembleRetriever  (vector + BM25)

class _PGVectorRetriever:
    def __init__(self, db: Session, job_embedding: Any, limit: int = 200):
        self.db = db
        self.job_embedding = job_embedding
        self.limit = limit

    def get_relevant_documents(self) -> list[tuple[UUID, float]]:
        print(f"\n[PGVECTOR] ---- START ----")
        print(f"[PGVECTOR] limit={self.limit} | emb_dim={self.job_embedding.embedding_dimension}")
        total_embs = self.db.query(CandidateEmbedding).count()
        print(f"[PGVECTOR] Total candidate embeddings: {total_embs}")
        if total_embs == 0:
            print("[PGVECTOR] WARNING: No candidate embeddings – vector search returns nothing!")
        stmt = (
            select(
                CandidateEmbedding.candidate_id,
                (1 - CandidateEmbedding.embedding.cosine_distance(
                    self.job_embedding.embedding
                )).label("score"),
            )
            .order_by(
                CandidateEmbedding.embedding.cosine_distance(self.job_embedding.embedding)
            )
            .limit(self.limit)
        )
        rows = self.db.execute(stmt).all()
        results = [(row[0], float(row[1])) for row in rows]
        print(f"[PGVECTOR] Returned {len(results)} candidates")
        if results:
            print(f"[PGVECTOR] Top-5: {[(str(c)[:8], round(s,4)) for c,s in results[:5]]}")
        print(f"[PGVECTOR] ---- END ----")
        return results


class _BM25Retriever:
    def __init__(self, db: Session, search_query: str, limit: int = 200):
        self.db = db
        self.search_query = search_query
        self.limit = limit

    def get_relevant_documents(self) -> list[tuple[UUID, float]]:
        print(f"\n[BM25] ---- START ----")
        print(f"[BM25] query='{self.search_query[:150]}' | limit={self.limit}")
        sql = text("""
            SELECT c.id,
                   ts_rank(
                       to_tsvector('english',
                           COALESCE(c.skills::text, '') || ' ' ||
                           COALESCE(c.tech_stack::text, '') || ' ' ||
                           COALESCE(c.sector_experience::text, '') || ' ' ||
                           COALESCE(c.raw_text, '')
                       ),
                       plainto_tsquery('english', :q)
                   ) AS bm25_score
            FROM candidates c
            ORDER BY bm25_score DESC
            LIMIT :lim
        """)
        rows = self.db.execute(sql, {"q": self.search_query, "lim": self.limit}).all()
        results = [(row[0], float(row[1])) for row in rows]
        print(f"[BM25] Returned {len(results)} candidates")
        if results:
            print(f"[BM25] Top-5: {[(str(c)[:8], round(s,4)) for c,s in results[:5]]}")
        print(f"[BM25] ---- END ----")
        return results


class EnsembleRetriever:
    def __init__(self, retrievers: list, weights: list[float]):
        assert len(retrievers) == len(weights)
        total = sum(weights)
        self.retrievers = retrievers
        self.weights = [w / total for w in weights]

    def get_relevant_documents(self, top_n: int = 100) -> list[tuple[UUID, float]]:
        print(f"\n[ENSEMBLE] ---- START ----  top_n={top_n}")
        all_scores: dict[UUID, float] = {}

        for i, (retriever, weight) in enumerate(zip(self.retrievers, self.weights)):
            rname = type(retriever).__name__
            docs = retriever.get_relevant_documents()
            print(f"[ENSEMBLE] [{i}] {rname} weight={round(weight,3)} → {len(docs)} docs")
            if not docs:
                continue
            max_score = max(score for _, score in docs) or 1.0
            for cid, raw_score in docs:
                norm = min(max(raw_score / max_score, 0.0), 1.0)
                all_scores[cid] = all_scores.get(cid, 0.0) + weight * norm

        ranked = sorted(all_scores.items(), key=lambda x: x[1], reverse=True)
        final = ranked[:top_n]
        print(f"[ENSEMBLE] Unique candidates={len(all_scores)} → returning top-{top_n}: {len(final)}")
        if final:
            print(f"[ENSEMBLE] Top-5: {[(str(c)[:8], round(s,4)) for c,s in final[:5]]}")
        print(f"[ENSEMBLE] ---- END ----")
        return final


def _build_bm25_query(job: Job) -> str:
    terms: set[str] = set()
    for field in [
        job.ai_required_skills, job.required_skills,
        job.ai_keywords, job.ai_must_have_keywords,
        job.ai_tools, job.ai_technologies,
    ]:
        if field:
            terms.update(t.lower().strip() for t in field if t and t.strip())
    return " ".join(terms) if terms else (job.job_title or "").lower()


def hybrid_search_candidates(
    db: Session,
    job_id: UUID,
    top_n: int = 100,
    vector_limit: int = 200,
    bm25_limit: int = 200,
    vector_weight: float = 0.7,
    bm25_weight: float = 0.3,
) -> list[tuple[UUID, float]]:
    print(f"\n{'='*60}")
    print(f"[HYBRID_SEARCH] START  job_id={job_id}  top_n={top_n}")
    logger.info("=== HYBRID SEARCH  job=%s  top_n=%d ===", job_id, top_n)

    job_embedding = db.query(JobEmbedding).filter(JobEmbedding.job_id == job_id).first()
    retrievers: list = []
    weights: list[float] = []

    if job_embedding and job_embedding.embedding is not None:
        retrievers.append(_PGVectorRetriever(db, job_embedding, limit=vector_limit))
        weights.append(vector_weight)
        print(f"[HYBRID_SEARCH] ✓ PGVector retriever added")
    else:
        print(f"[HYBRID_SEARCH] ✗ PGVector skipped (no embedding)")
        logger.warning("No job embedding for job %s", job_id)

    job = get_job_or_none(db, job_id)
    if job:
        q = _build_bm25_query(job)
        retrievers.append(_BM25Retriever(db, q, limit=bm25_limit))
        weights.append(bm25_weight)
        print(f"[HYBRID_SEARCH] ✓ BM25 retriever added  query='{q[:100]}'")
    else:
        print(f"[HYBRID_SEARCH] ✗ BM25 skipped (job not found)")

    if not retrievers:
        logger.warning("No retrievers available for job %s", job_id)
        return []

    ensemble = EnsembleRetriever(retrievers=retrievers, weights=weights)
    results = ensemble.get_relevant_documents(top_n=top_n)
    print(f"[HYBRID_SEARCH] → {len(results)} candidates  END")
    print(f"{'='*60}")
    return results


def vector_search_candidates(db: Session, job_id: UUID, limit: int = 200) -> list[tuple[UUID, float]]:
    job_embedding = db.query(JobEmbedding).filter(JobEmbedding.job_id == job_id).first()
    if job_embedding is None or job_embedding.embedding is None:
        return []
    return _PGVectorRetriever(db, job_embedding, limit=limit).get_relevant_documents()


def bm25_search_candidates(db: Session, job_id: UUID, limit: int = 200) -> list[tuple[UUID, float]]:
    job = get_job_or_none(db, job_id)
    if not job:
        return []
    return _BM25Retriever(db, _build_bm25_query(job), limit=limit).get_relevant_documents()


# Section 8  –  LLM Re-ranking  (weight 20%)

def llm_rerank_candidates(
    job: Job,
    match_results: list[MatchResult],
) -> dict[UUID, float]:
    """
    LLM contribution is 20% in final blend.
    project_impact_score and seniority_level used for explainability.
    Note: project_imp is now scored separately via score_project_impact (FIX P6).
    LLM here provides holistic fit + seniority context.
    """
    if not match_results:
        return {}

    print(f"\n[LLM_RERANK] ---- START ----  candidates={len(match_results)}")
    llm = get_llm(temperature=0.1)
    llm_scores: dict[UUID, float] = {}

    for result in match_results:
        candidate: Candidate = result.candidate
        if not candidate:
            print(f"[LLM_RERANK] WARNING: MatchResult has no candidate – skipping")
            continue

        prompt = f"""You are a senior technical recruiter evaluating candidates.

=== JOB ===
Title       : {job.job_title}
Description : {(job.job_description or '')[:1500]}
Required    : {', '.join((job.ai_required_skills or []) + (job.required_skills or []))}
Must-have   : {', '.join(job.ai_must_have_keywords or [])}
Experience  : {job.experience_min or 0}–{job.experience_max or '?'} years

=== CANDIDATE ===
Name        : {candidate.full_name or 'Unknown'}
Experience  : {candidate.total_experience_years or 'Unknown'} years
Skills      : {', '.join(candidate.skills or [])}
Tech stack  : {', '.join(candidate.tech_stack or [])}
Education   : {candidate.education_degree or ''} {candidate.education_field or ''}
Sector exp  : {', '.join(candidate.sector_experience or [])}

=== ATS PRE-SCORE ===
Overall ATS : {result.overall_score:.1f}
Matched     : {', '.join((result.matched_skills or [])[:10])}
Missing     : {', '.join((result.missing_skills or [])[:10])}

Evaluate and return ONLY valid JSON:
{{
  "llm_score": <0-100, overall fit>,
  "seniority_level": "<junior|mid|senior|lead|architect>",
  "seniority_confidence": <0.0-1.0>,
  "strengths": "<one concise sentence>",
  "weaknesses": "<one concise sentence>",
  "missing_skills": ["<skill1>", "<skill2>"]
}}
"""
        print(f"[LLM_RERANK] Evaluating: {candidate.full_name or 'Unknown'}")
        try:
            response = llm.invoke(prompt)
            raw = response.content if hasattr(response, "content") else str(response)
            parsed = parse_json_response(raw)
            score = max(0.0, min(100.0, float(parsed.get("llm_score", result.overall_score))))
            llm_scores[candidate.id] = round(score, 2)
            print(
                f"[LLM_RERANK] ✓ {candidate.full_name}: llm={score:.1f} "
                f"seniority={parsed.get('seniority_level','?')}"
            )
            logger.info(
                "LLM re-rank %s: llm=%.1f seniority=%s",
                candidate.full_name, score, parsed.get("seniority_level", "?"),
            )
        except Exception as exc:
            print(f"[LLM_RERANK] ERROR {candidate.full_name}: {exc} – fallback to ATS")
            llm_scores[candidate.id] = float(result.overall_score)

    print(f"[LLM_RERANK] ---- END ----  scored={len(llm_scores)}/{len(match_results)}")
    return llm_scores


# Section 9  –  Main matching pipeline

_ATS_WEIGHTS = {
    "must_have":   0.60,   # FIX P8: was 0.50
    "semantic":    0.15,   # FIX P8: was 0.20
    "project_imp": 0.10,   # FIX P8: was 0.15
    "domain":      0.10,   # same
    "education":   0.05,   # same
    "experience":  0.00,   # zeroed; redistributed when JD has range (see below)
}
_ATS_WEIGHTS_WITH_EXP = {
    "must_have":   0.60,   # FIX P8: was 0.50
    "semantic":    0.10,   # FIX P8: was 0.15
    "project_imp": 0.10,   # FIX P8: was 0.15
    "experience":  0.10,
    "domain":      0.05,
    "education":   0.05,
}

_FINAL_ATS_WEIGHT = 0.80
_FINAL_LLM_WEIGHT = 0.20


def run_matching(db: Session, job_id: UUID) -> list[MatchResult]:
    """
    Full pipeline v5:
        1. Hybrid retrieval (EnsembleRetriever)          → top-100
        2. Skill cluster dedup + RULE-BASED mandatory tier
           FIX P7: mandatory = required_skills ∪ ai_required_skills (no LLM)
                   preferred/bonus = LLM-classified from secondary pools only
        3. Mandatory gate  (only MANDATORY skills gate)
           FIX P1: arch_family (0.72) does NOT pass mandatory gate
        4. ATS scoring with v5 weight distribution
           FIX P8: must_have=60%, semantic=15%, project=10%, domain=10%, edu=5%
           FIX P6: project scored on complexity/scale/ownership/impact
        5. LLM re-ranking on top-20                      → 20% weight
        6. final = 0.80*ATS + 0.20*LLM
    """
    print(f"\n{'#'*60}")
    print(f"[RUN_MATCHING] ===== START v5 =====  job_id={job_id}")
    print(f"{'#'*60}")
    logger.info("=== RUN MATCHING v5  job=%s ===", job_id)

    job = get_job_or_none(db, job_id)
    if not job:
        print(f"[RUN_MATCHING] ERROR: Job {job_id} not found")
        return []

    print(f"[RUN_MATCHING] Job: '{job.job_title}' | exp_range={job.experience_min}-{job.experience_max}")

    has_experience_range = (
        job.experience_min is not None and float(job.experience_min or 0) > 0
    ) or job.experience_max is not None

    if has_experience_range:
        effective_weights = dict(_ATS_WEIGHTS_WITH_EXP)
        print(f"[RUN_MATCHING] FIX-P5: Experience range found → using experience bucket")
    else:
        effective_weights = dict(_ATS_WEIGHTS)
        print(f"[RUN_MATCHING] FIX-P5: No experience range in JD → experience_weight=0")

    print(f"[RUN_MATCHING] Effective ATS weights: {effective_weights}")

    # STEP 1: Hybrid retrieval
    print(f"\n[RUN_MATCHING] --- STEP 1: Hybrid Retrieval ---")
    hybrid_results = hybrid_search_candidates(db, job_id, top_n=100)
    if not hybrid_results:
        logger.warning("Hybrid search returned 0 candidates for job %s", job_id)
        return []

    top_ids = [cid for cid, _ in hybrid_results]
    hybrid_score_map = {cid: score for cid, score in hybrid_results}
    candidates = db.query(Candidate).filter(Candidate.id.in_(top_ids)).all()
    print(f"[RUN_MATCHING] Retrieved {len(candidates)} candidates")

    # STEP 2: Job keyword sets + tier classification
    print(f"\n[RUN_MATCHING] --- STEP 2: Keyword Sets + Tier Classification ---")
    job_secondary_skills = (
        (job.certifications or [])
        + (job.ai_nice_to_have_keywords or [])
        + (job.ai_tools or [])
        + (job.ai_technologies or [])
        + (job.ai_soft_skills or [])
    )
    job_keywords = normalize_keywords(list(set(
        (job.required_skills or [])
        + (job.ai_required_skills or [])
        + (job.ai_keywords or [])
        + (job.ai_must_have_keywords or [])
        + job_secondary_skills
        + [job.job_title, job.department, job.education_requirements]
    )))
    job_required_skills = list(set(
        (job.ai_required_skills or []) + (job.required_skills or [])
    ))
    job_tech_skills = job_required_skills
    certifications = job.certifications or []

    print(f"[RUN_MATCHING] job_keywords={len(job_keywords)}")

    if job.skill_tiers:
        weighted_must_have = job.skill_tiers
        print(f"[RUN_MATCHING] Using stored skill_tiers from job creation")
    else:
        weighted_must_have = get_weighted_must_have_skills(job)
        print(f"[RUN_MATCHING] No skill_tiers found, computing on the fly")
    mandatory_skills_list = [s for s, t in weighted_must_have.items() if t == "mandatory"]
    print(f"[RUN_MATCHING] weighted_must_have ({len(weighted_must_have)}): {weighted_must_have}")
    print(f"[RUN_MATCHING] mandatory_skills ({len(mandatory_skills_list)}): {mandatory_skills_list}")

    # STEP 3: ATS scoring
    results: list[MatchResult] = []
    print(f"\n[RUN_MATCHING] --- STEP 3: ATS Scoring ({len(candidates)} candidates) ---")

    for idx, candidate in enumerate(candidates):
        print(f"\n[ATS #{idx+1}/{len(candidates)}] {candidate.full_name or 'Unknown'} (id={str(candidate.id)[:8]}...)")

        candidate_keywords = normalize_keywords(
            (candidate.skills or [])
            + (candidate.tech_stack or [])
            + (candidate.sector_experience or [])
            + [candidate.education_degree, candidate.education_field, candidate.raw_text or ""]
        )
        matched = sorted(job_keywords & candidate_keywords)
        unmatched = sorted(job_keywords - candidate_keywords)

        # Skill proficiency / recency
        cand_parsed_meta = {}
        if hasattr(candidate, "extra_data") and candidate.extra_data:
            cand_parsed_meta = candidate.extra_data if isinstance(candidate.extra_data, dict) else {}
        skill_prof_map = build_skill_proficiency_map(cand_parsed_meta)

        # Score components
        score_exp    = score_experience(job, candidate) if has_experience_range else 0.0
        score_sec    = (keyword_overlap_score(job.ai_domain_experience, candidate.sector_experience or [])
                        if job.ai_domain_experience else 0.0)
        score_tech   = (fuzzy_skill_match_score(job_tech_skills, candidate.tech_stack or candidate.skills or [])
                        if job_tech_skills else 0.0)
        score_skill  = (fuzzy_skill_match_score(job_required_skills, candidate.skills or [])
                        if job_required_skills else 0.0)
        score_edu    = _score_education(job, candidate)
        score_certs  = (keyword_overlap_score(certifications, (candidate.skills or []) + (candidate.tech_stack or []))
                        if certifications else 0.0)
        score_other  = (keyword_overlap_score(job_secondary_skills,
                                              (candidate.skills or []) + (candidate.tech_stack or []))
                        if job_secondary_skills else 0.0)
        semantic_score = float(calculate_semantic_score(db, job, candidate))
        keyword_score  = float(len(matched) / max(len(job_keywords), 1) * 100)

        must_have_score, missing_mandatory_skills = score_must_have_skills(
            job, candidate, weighted_must_have, skill_prof_map
        )

        project_impact_score = score_project_impact(job, candidate)

        jskills_norm   = normalize_keywords(job_required_skills)
        cskills_norm   = normalize_keywords(candidate.skills or [])
        matched_skills = sorted(jskills_norm & cskills_norm)
        missing_skills = sorted(jskills_norm - cskills_norm)
        jtech_norm     = normalize_keywords(job_tech_skills)
        ctech_norm     = normalize_keywords(candidate.tech_stack or candidate.skills or [])
        matched_tech   = sorted(jtech_norm & ctech_norm)
        missing_tech   = sorted(jtech_norm - ctech_norm)

        print(
            f"[ATS #{idx+1}] must_have={must_have_score:.1f} semantic={semantic_score:.1f} "
            f"exp={score_exp:.1f} project_imp={project_impact_score:.1f}"
        )
        print(f"[ATS #{idx+1}] tech={score_tech:.1f} skill={score_skill:.1f} edu={score_edu:.1f} domain={score_sec:.1f}")

        # FIX P3: Mandatory gate (preferred/bonus do NOT gate)
        gate_failed, mandatory_ratio = apply_mandatory_gate(candidate, weighted_must_have, idx + 1)

        # FIX P8: ATS formula with v5 weights
        ats_score = float(round(
            must_have_score      * effective_weights["must_have"]
            + semantic_score     * effective_weights["semantic"]
            + score_exp          * effective_weights["experience"]
            + project_impact_score * effective_weights["project_imp"]
            + score_edu          * effective_weights["education"]
            + score_sec          * effective_weights["domain"],
            2,
        ))

        print(f"[ATS #{idx+1}] Raw ATS (before penalties) = {ats_score:.2f}")

        # Penalties
        if gate_failed:
            print(f"[ATS #{idx+1}] GATE PENALTY: {ats_score:.2f} × 0.20 = {ats_score*0.2:.2f}")
            logger.info("Gate penalty: ATS %.2f → %.2f for %s", ats_score, ats_score * 0.2, candidate.full_name)
            ats_score = round(ats_score * 0.2, 2)
        elif weighted_must_have and must_have_score < 30:
            print(f"[ATS #{idx+1}] SOFT PENALTY: must_have<30% → {ats_score:.2f} × 0.70")
            ats_score = round(ats_score * 0.7, 2)

        print(f"[ATS #{idx+1}] FINAL ATS = {ats_score:.2f}")
        logger.info(
            "ATS v5 | must=%.1f sem=%.1f exp=%.1f proj=%.1f edu=%.1f dom=%.1f gate=%s → %.2f",
            must_have_score, semantic_score, score_exp, project_impact_score,
            score_edu, score_sec, gate_failed, ats_score,
        )

        # Upsert MatchResult
        result = (
            db.query(MatchResult)
            .filter(
                MatchResult.job_posting_id == job_id,
                MatchResult.candidate_id == candidate.id,
            )
            .first()
        )
        if not result:
            result = MatchResult(job_posting_id=job_id, candidate_id=candidate.id)
            db.add(result)

        result.overall_score       = ats_score
        result.score_experience    = score_exp
        result.score_sector        = score_sec
        result.score_tech_stack    = score_tech
        result.score_skill         = score_skill
        result.score_education     = score_edu
        result.score_other_skills  = score_other
        result.matched_keywords    = matched
        result.unmatched_keywords  = unmatched
        result.matched_skills      = matched_skills
        result.missing_skills      = missing_skills
        result.matched_tech_stack  = matched_tech
        result.missing_tech_stack  = missing_tech
        result.bm25_score          = round(keyword_score, 4)
        result.semantic_score      = semantic_score
        result.ai_summary          = build_match_summary(candidate, matched, unmatched, ats_score)
        if hasattr(result, "missing_critical_skills"):
            result.missing_critical_skills = missing_mandatory_skills
        if hasattr(result, "project_impact_score"):
            result.project_impact_score = project_impact_score

        results.append(result)

    db.commit()

    # STEP 4: LLM re-ranking  (top-20)
    print(f"\n[RUN_MATCHING] --- STEP 4: LLM Re-ranking (top-20) ---")
    results_sorted = sorted(results, key=lambda r: r.overall_score, reverse=True)
    top_20 = results_sorted[:20]

    for r in top_20:
        db.refresh(r)

    llm_scores = llm_rerank_candidates(job, top_20)
    print(f"[RUN_MATCHING] LLM scored {len(llm_scores)}/{len(top_20)}")

    # STEP 5: Blend ATS + LLM  (80/20)
    print(f"\n[RUN_MATCHING] --- STEP 5: Blend ATS ({int(_FINAL_ATS_WEIGHT*100)}%) + LLM ({int(_FINAL_LLM_WEIGHT*100)}%) ---")
    for result in results:
        ats = float(result.overall_score)
        llm = float(llm_scores.get(result.candidate_id, ats))
        final = round(_FINAL_ATS_WEIGHT * ats + _FINAL_LLM_WEIGHT * llm, 2)
        name = getattr(result.candidate, "full_name", None) or str(result.candidate_id)[:8]
        print(f"[BLEND] {name}: ATS={ats:.2f}  LLM={llm:.2f}  → FINAL={final:.2f}")
        logger.info("Final %s: ATS=%.2f LLM=%.2f → %.2f", result.candidate_id, ats, llm, final)
        result.overall_score = final

    db.commit()

    # STEP 6: Rank and return
    print(f"\n[RUN_MATCHING] --- STEP 6: Final Ranking ---")
    ranked = (
        db.query(MatchResult)
        .options(joinedload(MatchResult.candidate))
        .filter(MatchResult.job_posting_id == job_id)
        .order_by(MatchResult.overall_score.desc())
        .all()
    )
    for pos, r in enumerate(ranked, start=1):
        r.rank_position = pos
    db.commit()

    print(f"\n[RUN_MATCHING] ===== FINAL TOP-10 =====")
    for r in ranked[:10]:
        print(f"  #{r.rank_position:>2}  {(getattr(r.candidate,'full_name','Unknown') or 'Unknown'):<30}  score={float(r.overall_score):.2f}")
    print(f"[RUN_MATCHING] ===== END =====  total={len(ranked)}")
    print(f"{'#'*60}\n")
    logger.info("Matching v5 complete: %d ranked for job %s", len(ranked), job_id)
    return ranked


# Section 10  –  Scoring helpers

def _score_education(job: Job, candidate: Candidate) -> float:
    """Education scored 0/50/100; capped at 5% by ATS weight."""
    req = (job.education_requirements or "").strip()
    deg = (candidate.education_degree or "").strip()
    if not req:
        return 50.0 if deg else 0.0
    if deg and deg.lower() in req.lower():
        return 100.0
    return 0.0


def score_experience(job: Job, candidate: Candidate) -> float:
    """Called only when has_experience_range is True."""
    years   = float(candidate.total_experience_years or 0)
    exp_min = float(job.experience_min or 0)
    exp_max = float(job.experience_max) if job.experience_max is not None else None

    if exp_min == 0:
        return min(round(years / 10 * 100, 2), 100.0) if years > 0 else 0.0

    if years >= exp_min and (exp_max is None or years <= exp_max):
        return 100.0

    if exp_max is not None and years > exp_max:
        ratio = years / exp_max
        if ratio > 2:   return 70.0
        if ratio > 1.5: return 80.0
        return 85.0

    if years < exp_min:
        return round(years / max(exp_min, 1) * 100, 2)

    return 100.0


def calculate_semantic_score(db: Session, job: Job, candidate: Candidate) -> float:
    job_emb  = db.query(JobEmbedding).filter(JobEmbedding.job_id == job.job_id).first()
    cand_emb = db.query(CandidateEmbedding).filter(CandidateEmbedding.candidate_id == candidate.id).first()
    if not job_emb or not cand_emb:
        return 0.0
    jv, cv = job_emb.embedding, cand_emb.embedding
    if jv is None or cv is None:
        return 0.0
    sim   = cosine_similarity(list(jv), list(cv))
    score = float(round(max(min(sim, 1.0), -1.0) * 50 + 50, 4))
    print(f"  [SEMANTIC] cosine={sim:.4f} → {score:.4f}")
    return score


def cosine_similarity(left: list[float], right: list[float]) -> float:
    if not left or not right or len(left) != len(right):
        return 0.0
    dot   = sum(float(a) * float(b) for a, b in zip(left, right))
    lnorm = math.sqrt(sum(float(a) ** 2 for a in left))
    rnorm = math.sqrt(sum(float(b) ** 2 for b in right))
    if lnorm == 0 or rnorm == 0:
        return 0.0
    return dot / (lnorm * rnorm)


def _simple_token_ratio(a: str, b: str) -> int:
    ta = set(a.lower().split())
    tb = set(b.lower().split())
    if not ta and not tb: return 100
    if not ta or not tb:  return 0
    return int(len(ta & tb) / len(ta | tb) * 100)


def fuzzy_skill_match_score(required: list[str], actual: list[str]) -> float:
    if not required or not actual:
        return 0.0
    try:
        from thefuzz import fuzz
        _ratio = lambda a, b: fuzz.token_set_ratio(a.lower(), b.lower())
    except ImportError:
        _ratio = _simple_token_ratio
    total = 0.0
    for req in required:
        best = max(_ratio(req, act) for act in actual)
        total += best if best >= 70 else best * 0.5
    return round(total / len(required), 2)


def keyword_overlap_score(required: list[str], actual: list[str]) -> float:
    if not required:
        return 100.0
    act_flat: set[str] = normalize_keywords([a for a in actual if a])
    matched = 0
    for req in required:
        rf = next(iter(normalize_keywords([req])), "")
        if rf and rf in act_flat:
            matched += 1
            continue
        for act in actual:
            if rf and rf in act.lower():
                matched += 1
                break
    return round(matched / len(required) * 100, 2)


def build_match_summary(
    candidate: Candidate,
    matched: list[str],
    unmatched: list[str],
    overall: float,
) -> str:
    return (
        f"{candidate.full_name or 'Candidate'} scored {overall:.2f}. "
        f"Matched: {', '.join(matched[:8]) or 'none'}. "
        f"Missing: {', '.join(unmatched[:8]) or 'none'}."
    )


# Section 11  –  Resume processing

def get_job_or_none(db: Session, job_id: UUID) -> Job | None:
    return db.query(Job).filter(Job.job_id == job_id).first()


def get_resume_files(db: Session, job_id: UUID) -> list[ResumeFile]:
    return (
        db.query(ResumeFile)
        .filter(ResumeFile.job_posting_id == job_id)
        .order_by(ResumeFile.created_at.desc())
        .all()
    )


async def save_resume_files(
    db: Session,
    job_id: UUID,
    files: list[UploadFile],
) -> list[ResumeFile]:
    upload_dir = Path(settings.RESUME_UPLOAD_DIR) / str(job_id)
    upload_dir.mkdir(parents=True, exist_ok=True)

    saved: list[tuple[ResumeFile, bool]] = []
    for upload in files:
        original_filename = Path(upload.filename or "resume").name
        extension         = Path(original_filename).suffix.lower()
        content           = await upload.read()
        file_hash         = hashlib.md5(content).hexdigest()
        storage_name      = f"{uuid4()}_{original_filename}"
        storage_path      = upload_dir / storage_name

        with storage_path.open("wb") as buf:
            buf.write(content)

        is_dup = (
            db.query(ResumeFile)
            .filter(
                ResumeFile.job_posting_id == job_id,
                ResumeFile.file_hash_md5 == file_hash,
            )
            .first() is not None
        )

        rf = ResumeFile(
            job_posting_id=job_id,
            original_filename=original_filename,
            storage_path=str(storage_path),
            file_size_bytes=len(content),
            file_hash_md5=file_hash,
        )
        if is_dup:
            rf.validation_status = ResumeValidationStatus.INVALID
            rf.processing_status = ResumeProcessingStatus.FAILED
            rf.rejection_reason  = "Duplicate resume hash for this job"
        elif extension not in ALLOWED_RESUME_EXTENSIONS:
            rf.validation_status = ResumeValidationStatus.INVALID
            rf.processing_status = ResumeProcessingStatus.FAILED
            rf.rejection_reason  = f"Unsupported file type: {extension}. Upload PDF, DOCX, or TXT."
        db.add(rf)
        saved.append((rf, is_dup or extension not in ALLOWED_RESUME_EXTENSIONS))

    db.commit()

    result: list[ResumeFile] = []
    for rf, skip in saved:
        db.refresh(rf)
        if not skip:
            from app.tasks.resume_tasks import process_single_resume_task
            try:
                process_single_resume_task.apply_async(
                    args=[str(rf.id)],
                    queue="resume_processing",
                )
            except Exception as exc:
                logger.exception("Failed to enqueue resume task")
                mark_resume_failed(db, rf, f"Enqueue failed: {exc}", remove_local=False)
        result.append(rf)

    return result


def delete_resume_file(db: Session, job_id: UUID, resume_file_id: UUID) -> bool:
    rf = (
        db.query(ResumeFile)
        .filter(ResumeFile.job_posting_id == job_id, ResumeFile.id == resume_file_id)
        .first()
    )
    if not rf:
        return False
    p = Path(rf.storage_path)
    if p.exists():
        p.unlink()
    db.delete(rf)
    db.commit()
    return True


def process_single_resume(db: Session, resume_file: ResumeFile) -> None:
    resume_file.processing_status = ResumeProcessingStatus.PROCESSING
    db.commit()

    try:
        raw_text  = extract_resume_text(resume_file.storage_path)
        valid, reason = validate_resume_text(raw_text)
        if not valid:
            mark_resume_failed(db, resume_file, reason, remove_local=True)
            return

        parsed        = extract_candidate_details(raw_text)
        emb_vec, emb_model = generate_candidate_embedding(raw_text, parsed)
        emb_dim       = len(emb_vec)

        # Check for duplicate resume by file hash for this job
        duplicate = (
            db.query(ResumeFile)
            .join(Candidate, Candidate.resume_file_id == ResumeFile.id)
            .filter(
                ResumeFile.job_posting_id == resume_file.job_posting_id,
                ResumeFile.file_hash_md5  == resume_file.file_hash_md5,
                ResumeFile.id             != resume_file.id,
            )
            .first()
        )
        
        # If duplicate found, mark this resume file as invalid and skip processing
        if duplicate:
            resume_file.validation_status = ResumeValidationStatus.INVALID
            resume_file.processing_status = ResumeProcessingStatus.FAILED
            resume_file.rejection_reason = "Duplicate resume hash for this job"
            db.commit()
            return {"status": "duplicate", "message": "Duplicate resume detected"}

        candidate = (
            db.query(Candidate)
            .filter(Candidate.resume_file_id == resume_file.id)
            .first()
        )
        if not candidate:
            candidate = Candidate(resume_file_id=resume_file.id)
            db.add(candidate)

        candidate.full_name              = parsed.get("full_name")
        candidate.email                  = parsed.get("email")
        candidate.phone                  = parsed.get("phone")
        candidate.total_experience_years = parse_optional_float(parsed.get("total_experience_years"))
        candidate.education_degree       = parsed.get("education_degree")
        candidate.education_field        = parsed.get("education_field")
        candidate.skills                 = normalize_list(parsed.get("skills"))
        candidate.tech_stack             = normalize_list(parsed.get("tech_stack"))
        candidate.sector_experience      = normalize_list(parsed.get("sector_experience"))
        candidate.raw_text               = raw_text
        candidate.is_duplicate           = duplicate
        
        # Extract projects from raw text and store as JSON
        projects = extract_projects_from_raw_text(raw_text)
        candidate.projects_json = {
            "projects": projects
        } if projects else None
        
        # Extract experience sectors and store as JSON
        candidate.experience_json = {
            "sectors": normalize_list(parsed.get("sector_experience")),
            "total_years": parse_optional_float(parsed.get("total_experience_years"))
        } if parsed.get("sector_experience") or parsed.get("total_experience_years") else None
        
        if hasattr(candidate, "extra_data") and parsed.get("skill_proficiency"):
            candidate.extra_data = {
                **(candidate.extra_data or {}),
                "skill_proficiency": parsed["skill_proficiency"],
            }

        resume_file.validation_status = ResumeValidationStatus.VALID
        resume_file.rejection_reason  = None
        resume_file.processing_status = ResumeProcessingStatus.COMPLETED
        db.commit()
        db.refresh(candidate)

        content_hash = hashlib.md5(raw_text.encode()).hexdigest()
        embedding = (
            db.query(CandidateEmbedding)
            .filter(CandidateEmbedding.candidate_id == candidate.id)
            .first()
        )
        if not embedding:
            embedding = CandidateEmbedding(
                candidate_id=candidate.id,
                embedding=emb_vec,
                embedding_model=emb_model,
                embedding_dimension=emb_dim,
                content_hash=content_hash,
                generated_at=datetime.utcnow(),
            )
            db.add(embedding)
        else:
            embedding.embedding           = emb_vec
            embedding.embedding_model     = emb_model
            embedding.embedding_dimension = emb_dim
            embedding.content_hash        = content_hash
            embedding.generated_at        = datetime.utcnow()

        db.commit()
        db.refresh(embedding)
        candidate.embedding_id = str(embedding.embedding_id)
        db.commit()

    except Exception as exc:
        logger.exception("Resume processing failed")
        mark_resume_failed(db, resume_file, str(exc), remove_local=False)
        raise


def mark_resume_failed(
    db: Session,
    resume_file: ResumeFile,
    reason: str,
    remove_local: bool,
) -> None:
    resume_file.validation_status = ResumeValidationStatus.INVALID
    resume_file.rejection_reason  = reason[:500]
    resume_file.processing_status = ResumeProcessingStatus.FAILED
    db.commit()
    if remove_local:
        p = Path(resume_file.storage_path)
        if p.exists():
            p.unlink()


def extract_resume_text(storage_path: str) -> str:
    path = Path(storage_path)
    if not path.exists():
        raise FileNotFoundError("Uploaded file missing from local storage")

    ext = path.suffix.lower()
    if ext not in ALLOWED_RESUME_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    if ext == ".pdf":
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages).strip()

    if ext == ".docx":
        from docx import Document
        doc   = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()

    return path.read_text(encoding="utf-8", errors="ignore").strip()


def is_financial_document(text: str) -> bool:
    lowered = text.lower()
    hits = sum(1 for p in FINANCIAL_DOC_PATTERNS if re.search(p, lowered))
    if hits >= 2:
        return True
    if "invoice" in lowered:
        return not any(re.search(p, lowered) for p in RESUME_PROJECT_CONTEXT)
    return False


def validate_resume_text(text: str) -> tuple[bool, str | None]:
    if len(text.strip()) < 200:
        return False, "File does not contain enough readable resume text"
    try:
        llm = get_llm(temperature=0.1)
        prompt = f"""Analyze this text and determine if it is a resume/CV.
Return ONLY valid JSON:
{{"is_resume": <true|false>, "reason": "<brief explanation>"}}
Text:
{text[:8000]}
"""
        response = llm.invoke(prompt)
        raw      = response.content if hasattr(response, "content") else str(response)
        parsed   = parse_json_response(raw)
        if not parsed.get("is_resume", False):
            return False, f"Not a resume: {parsed.get('reason', 'unknown')}"
        return True, None
    except Exception as exc:
        logger.warning("AI validation failed, using keyword fallback: %s", exc)
        lowered = text.lower()
        signals = sum(1 for kw in RESUME_KEYWORDS if kw in lowered)
        if extract_email(text) or extract_phone(text):
            signals += 1
        if extract_years_experience(text) is not None:
            signals += 1
        if signals < 2:
            return False, "Content does not look like a valid resume"
        return True, None


def extract_candidate_details(text: str) -> dict[str, Any]:
    fallback = fallback_candidate_extract(text)
    try:
        llm = get_llm(temperature=0.1)
        prompt = f"""Extract candidate details from this resume.

IMPORTANT: Extract skills from ALL sections — Skills, Experience, Projects, Achievements.
Include skills implied by work descriptions even if not listed in Skills section.
Examples:
  - "Built distributed Kafka pipeline" → add "kafka", "distributed systems"
  - "Led team of 8 engineers" → add "leadership", "team management"
  - "Designed REST APIs using Spring Boot" → add "spring boot", "rest", "java", "microservices"
  - "Designed microservices with Spring Cloud" → add "enterprise application architecture", "microservices"
  - "Built serverless functions on AWS Lambda" → add "serverless", "aws"

Return ONLY valid JSON (null for unknown scalars, [] for unknown lists):
{{
  "full_name": "<string>",
  "email": "<string>",
  "phone": "<string>",
  "total_experience_years": <number or null>,
  "education_degree": "<string>",
  "education_field": "<string>",
  "skills": ["<all skills including implied ones>"],
  "tech_stack": ["<technologies actually used from all sections>"],
  "sector_experience": ["<domains/industries>"],
  "skills_from_experience": ["<skills inferred only from job descriptions>"],
  "skill_proficiency": [
    {{"skill": "<skill name>", "years": <number or null>, "last_used_year": <4-digit year or null>, "confidence": <0.0-1.0>}}
  ]
}}

Resume:
{text[:12000]}
"""
        response = llm.invoke(prompt)
        raw      = response.content if hasattr(response, "content") else str(response)
        parsed   = parse_json_response(raw)

        implied       = parsed.pop("skills_from_experience", []) or []
        existing      = normalize_list(parsed.get("skills"))
        parsed["skills"] = list(dict.fromkeys(existing + normalize_list(implied)))

        return {**fallback, **{k: v for k, v in parsed.items() if v not in (None, "", [])}}
    except Exception as exc:
        logger.warning("LLM extraction failed, fallback: %s", exc)
        return fallback


def fallback_candidate_extract(text: str) -> dict[str, Any]:
    lines  = [ln.strip() for ln in text.splitlines() if ln.strip()]
    name   = next((ln for ln in lines[:10]
                   if not extract_email(ln) and not extract_phone(ln)), None)
    skills = extract_section_terms(text, ["skills", "technical skills", "technologies"])
    edu    = extract_section_terms(text, ["education", "academic"])
    return {
        "full_name":              name[:200] if name else None,
        "email":                  extract_email(text),
        "phone":                  extract_phone(text),
        "total_experience_years": extract_years_experience(text),
        "education_degree":       edu[0][:150] if edu else None,
        "education_field":        None,
        "skills":                 skills,
        "tech_stack":             skills,
        "sector_experience":      [],
    }


def generate_candidate_embedding(
    text: str, parsed: dict[str, Any]
) -> tuple[list[float], str]:
    embeddings = get_embeddings()
    emb_text   = build_candidate_embedding_text(text, parsed)
    vector     = embeddings.embed_query(emb_text)
    return normalize_embedding_vector(vector), resolve_embedding_model_name()


def build_candidate_embedding_text(text: str, parsed: dict[str, Any]) -> str:
    return (
        f"Skills:\n{', '.join(normalize_list(parsed.get('skills')))}\n\n"
        f"Tech Stack:\n{', '.join(normalize_list(parsed.get('tech_stack')))}\n\n"
        f"Domain:\n{', '.join(normalize_list(parsed.get('sector_experience')))}\n\n"
        f"Education:\n{parsed.get('education_degree') or ''} "
        f"{parsed.get('education_field') or ''}\n\n"
        f"Resume:\n{text[:3000]}"
    ).strip()


def resolve_embedding_model_name() -> str:
    provider = (settings.EMBEDDING_PROVIDER or "ollama").lower()
    if provider == "openai":
        return settings.OPENAI_EMBEDDING_MODEL
    if provider == "gemini":
        return settings.HUGGINGFACE_EMBEDDING_MODEL
    return (
        settings.AI_EMBEDDING_MODEL
        or settings.OLLAMA_EMBEDDING_MODEL
        or settings.OLLAMA_MODEL
        or "ollama"
    )


def get_processing_summary(db: Session, job_id: UUID) -> dict[str, Any]:
    db.expire_all()
    rows = (
        db.query(ResumeFile)
        .options(joinedload(ResumeFile.candidate))
        .filter(ResumeFile.job_posting_id == job_id)
        .order_by(ResumeFile.created_at.desc())
        .all()
    )
    def _count(status: ResumeProcessingStatus) -> int:
        return sum(1 for r in rows if r.processing_status == status)

    return {
        "total":      len(rows),
        "completed":  _count(ResumeProcessingStatus.COMPLETED),
        "failed":     _count(ResumeProcessingStatus.FAILED),
        "pending":    _count(ResumeProcessingStatus.PENDING),
        "processing": _count(ResumeProcessingStatus.PROCESSING),
        "logs":       [{"resume_file": r, "candidate": r.candidate} for r in rows],
    }


def get_candidates_for_job(db: Session, job_id: UUID) -> list[Candidate]:
    return (
        db.query(Candidate)
        .join(ResumeFile, Candidate.resume_file_id == ResumeFile.id)
        .filter(ResumeFile.job_posting_id == job_id)
        .order_by(Candidate.created_at.desc())
        .all()
    )


def extract_projects_from_raw_text(raw_text: str | None) -> list[dict]:
    """Extract projects from candidate's raw text using simple pattern matching."""
    if not raw_text:
        return []
    
    import re
    
    projects = []
    lines = raw_text.split('\n')
    
    current_project = None
    in_project_section = False
    
    for line in lines:
        line = line.strip()
        
        # Detect project section
        if 'project' in line.lower() and ':' in line:
            in_project_section = True
            continue
        
        # Detect project name (numbered or standalone)
        if in_project_section and (line.startswith('Project') or re.match(r'^\d+\.', line)):
            if current_project:
                projects.append(current_project)
            current_project = {
                "name": line.replace('Project', '').replace(':', '').strip(),
                "description": "",
                "technologies": []
            }
        elif current_project:
            # Extract description
            if line.lower().startswith('description'):
                current_project["description"] = line.replace('Description:', '').replace('description:', '').strip()
            # Extract technologies
            elif 'technologies' in line.lower() or 'tech stack' in line.lower():
                tech_line = line.replace('Technologies used:', '').replace('technologies used:', '').replace('Technologies:', '').replace('technologies:', '').strip()
                current_project["technologies"] = [t.strip() for t in tech_line.split(',') if t.strip()]
    
    if current_project:
        projects.append(current_project)
    
    return projects


def get_match_results(db: Session, job_id: UUID, limit: int = 100) -> list[dict]:
    from app.schemas.resume import ScoreBreakdown
    
    results = (
        db.query(MatchResult)
        .options(joinedload(MatchResult.candidate))
        .join(Candidate, MatchResult.candidate_id == Candidate.id)
        .join(ResumeFile, Candidate.resume_file_id == ResumeFile.id)
        .filter(
            MatchResult.job_posting_id == job_id,
            Candidate.is_duplicate == False
        )
        .order_by(
            MatchResult.rank_position.asc().nullslast(),
            MatchResult.overall_score.desc(),
        )
        .limit(limit)
        .all()
    )
    
    # Deduplicate results by candidate full_name to avoid showing same person twice
    seen_names = set()
    deduplicated_results = []
    for r in results:
        if r.candidate and r.candidate.full_name:
            if r.candidate.full_name not in seen_names:
                seen_names.add(r.candidate.full_name)
                deduplicated_results.append(r)
        else:
            deduplicated_results.append(r)
    
    results = deduplicated_results
    
    # Transform to new response structure with score breakdown
    response_data = []
    for r in results:
        # Calculate keyword score from matched_keywords
        keyword_score = None
        if r.matched_keywords is not None and r.matched_keywords:
            # This is a rough calculation - adjust as needed
            keyword_score = len(r.matched_keywords) * 10  # Simple calculation
        
        score_breakdown = ScoreBreakdown(
            semantic_score=float(r.semantic_score) if r.semantic_score else None,
            bm25_score=float(r.bm25_score) if r.bm25_score else None,
            keyword_score=keyword_score,
            skill_score=float(r.score_skill) if r.score_skill else None,
            tech_stack_score=float(r.score_tech_stack) if r.score_tech_stack else None,
            experience_score=float(r.score_experience) if r.score_experience else None,
            education_score=float(r.score_education) if r.score_education else None,
            sector_score=float(r.score_sector) if r.score_sector else None,
            other_skills_score=float(r.score_other_skills) if r.score_other_skills else None,
        )
        
        # Use projects from JSON field if available, otherwise extract from raw text
        # if r.candidate and r.candidate.projects_json:
        #     projects = r.candidate.projects_json.get("projects", [])
        # else:
        #     projects = extract_projects_from_raw_text(r.candidate.raw_text if r.candidate else None)
        
        # Build candidate dict with projects
        candidate_dict = None
        if r.candidate:
            candidate_dict = {
                "full_name": r.candidate.full_name,
                "email": r.candidate.email,
                "phone": r.candidate.phone,
                "total_experience_years": (
                    float(r.candidate.total_experience_years)
                    if r.candidate.total_experience_years
                    else None
                ),
                "education_degree": r.candidate.education_degree,
                "skills": r.candidate.skills or [],
                "tech_stack": r.candidate.tech_stack or [],
                "sector_experience": r.candidate.sector_experience or [],
                "raw_text": r.candidate.raw_text,
            }
        response_data.append({
            "overall_score": float(r.overall_score),
            "rank_position": r.rank_position,

            "candidate": candidate_dict,

            "matched_skills": (
                r.matched_skills
                if r.matched_skills
                else r.matched_tech_stack or []
            ),

            "score_breakdown": {
                "semantic_score": score_breakdown.semantic_score,
                "bm25_score": score_breakdown.bm25_score,
                "keyword_score": score_breakdown.keyword_score,
                "skill_score": score_breakdown.skill_score,
                "tech_stack_score": score_breakdown.tech_stack_score,
                "experience_score": score_breakdown.experience_score,
                "education_score": score_breakdown.education_score,
                "sector_score": score_breakdown.sector_score,
                "other_skills_score": score_breakdown.other_skills_score,
            }
        })
    
    return response_data


def parse_optional_float(value: Any) -> float | None:
    if value in (None, ""):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    m = re.search(r"\d+(?:\.\d+)?", str(value))
    return float(m.group(0)) if m else None


def normalize_embedding_vector(vector: Any) -> list[float]:
    return [float(v) for v in vector]


def normalize_list(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(i).strip() for i in value if str(i).strip()]
    if isinstance(value, str):
        return [i.strip() for i in re.split(r"[,;\n]", value) if i.strip()]
    return []


def extract_email(text: str) -> str | None:
    m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    return m.group(0) if m else None


def extract_phone(text: str) -> str | None:
    m = re.search(r"(\+?\d[\d\s().-]{8,}\d)", text)
    return m.group(1).strip() if m else None


def extract_years_experience(text: str) -> float | None:
    matches = re.findall(
        r"(\d+(?:\.\d+)?)\+?\s*(?:years|yrs)\s+(?:of\s+)?experience",
        text, re.IGNORECASE,
    )
    return max(float(m) for m in matches) if matches else None


def extract_section_terms(text: str, headings: list[str]) -> list[str]:
    lowered = text.lower()
    for heading in headings:
        start = lowered.find(heading)
        if start == -1:
            continue
        snippet = text[start: start + 800]
        terms   = re.split(r"[,|;\n]", snippet)
        return [t.strip(" -:\t") for t in terms[1:25] if 1 < len(t.strip()) < 80]
    return []
